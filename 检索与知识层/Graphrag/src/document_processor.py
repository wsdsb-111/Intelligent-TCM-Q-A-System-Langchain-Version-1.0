"""
中医文档处理引擎
支持多种文档格式的中医文本提取和处理
"""

import os
import re
from pathlib import Path
from typing import Optional, Dict, Any
import logging

from models import ProcessedDocument
from exceptions import DocumentProcessingError, FileProcessingError


class DocumentProcessor:
    """中医文档处理器"""
    
    def __init__(self, config=None):
        self.config = config
        self.logger = logging.getLogger(f"{__name__}.DocumentProcessor")
        self.supported_formats = {
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx
        }
    
    def process_document(self, file_path: str) -> ProcessedDocument:
        """
        处理文档并返回标准化的ProcessedDocument对象
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            ProcessedDocument: 处理后的文档对象
            
        Raises:
            DocumentProcessingError: 文档处理失败
            FileProcessingError: 文件读取失败
        """
        try:
            file_path = Path(file_path)
            
            # 检查文件是否存在
            if not file_path.exists():
                raise FileProcessingError(f"文件不存在: {file_path}")
            
            # 检查文件格式是否支持
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_formats:
                raise DocumentProcessingError(
                    f"不支持的文件格式: {file_extension}",
                    error_code="UNSUPPORTED_FORMAT"
                )
            
            self.logger.info(f"开始处理文档: {file_path}")
            
            # 提取文本内容
            processor_func = self.supported_formats[file_extension]
            content = processor_func(file_path)
            
            # 清理和标准化文本
            cleaned_content = self.clean_and_normalize(content)
            
            # 提取元数据
            metadata = self._extract_metadata(file_path, content)
            
            # 创建ProcessedDocument对象
            document = ProcessedDocument(
                title=file_path.stem,
                content=cleaned_content,
                file_path=str(file_path),
                file_type=file_extension,
                metadata=metadata
            )
            
            self.logger.info(f"文档处理完成: {file_path}, 内容长度: {len(cleaned_content)}")
            return document
            
        except Exception as e:
            self.logger.error(f"文档处理失败: {file_path}, 错误: {e}")
            if isinstance(e, (DocumentProcessingError, FileProcessingError)):
                raise
            raise DocumentProcessingError(f"文档处理失败: {e}")
    
    def _process_text(self, file_path: Path) -> str:
        """处理纯文本文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    self.logger.debug(f"成功使用编码 {encoding} 读取文件: {file_path}")
                    return content
                except UnicodeDecodeError:
                    continue
            
            raise DocumentProcessingError(f"无法解码文件: {file_path}")
            
        except Exception as e:
            raise DocumentProcessingError(f"读取文本文件失败: {e}")
    
    def _process_markdown(self, file_path: Path) -> str:
        """处理Markdown文件"""
        try:
            # Markdown文件本质上是文本文件
            content = self._process_text(file_path)
            
            # 可以在这里添加Markdown特定的处理逻辑
            # 例如：移除Markdown标记，保留纯文本
            # 暂时保持原样，后续可以集成markdown库进行解析
            
            return content
            
        except Exception as e:
            raise DocumentProcessingError(f"处理Markdown文件失败: {e}")
    
    def _process_pdf(self, file_path: Path) -> str:
        """处理PDF文件"""
        try:
            # 尝试导入PDF处理库
            try:
                import PyPDF2
            except ImportError:
                raise DocumentProcessingError(
                    "缺少PyPDF2库，请安装: pip install PyPDF2",
                    error_code="MISSING_DEPENDENCY"
                )
            
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 提取所有页面的文本
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    content += page_text + "\n"
            
            if not content.strip():
                self.logger.warning(f"PDF文件可能是扫描版或无文本内容: {file_path}")
                return ""
            
            return content
            
        except Exception as e:
            raise DocumentProcessingError(f"处理PDF文件失败: {e}")
    
    def _process_docx(self, file_path: Path) -> str:
        """处理Word文档"""
        try:
            # 尝试导入Word处理库
            try:
                from docx import Document
            except ImportError:
                raise DocumentProcessingError(
                    "缺少python-docx库，请安装: pip install python-docx",
                    error_code="MISSING_DEPENDENCY"
                )
            
            doc = Document(file_path)
            content = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            return content
            
        except Exception as e:
            raise DocumentProcessingError(f"处理Word文档失败: {e}")
    
    def clean_and_normalize(self, text: str) -> str:
        """
        清理和标准化文本
        
        Args:
            text: 原始文本
            
        Returns:
            str: 清理后的文本
        """
        if not text:
            return ""
        
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除行首行尾空白
        text = text.strip()
        
        # 标准化换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # 移除多余的换行符
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 移除特殊字符（保留中文、英文、数字、常用标点）
        text = re.sub(r'[^\u4e00-\u9fff\w\s.,!?;:()[]{}""''""—–-]', '', text)
        
        # 处理中文标点
        chinese_punctuation = {
            '，': ',', '。': '.', '！': '!', '？': '?', 
            '；': ';', '：': ':', '（': '(', '）': ')',
            '【': '[', '】': ']', '《': '<', '》': '>',
            '"': '"', '"': '"', ''': "'", ''': "'"
        }
        
        for chinese, english in chinese_punctuation.items():
            text = text.replace(chinese, english)
        
        return text
    
    def _extract_metadata(self, file_path: Path, content: str) -> Dict[str, Any]:
        """
        提取文档元数据
        
        Args:
            file_path: 文件路径
            content: 文档内容
            
        Returns:
            Dict[str, Any]: 元数据字典
        """
        metadata = {}
        
        try:
            # 文件基本信息
            stat = file_path.stat()
            metadata.update({
                'file_size': stat.st_size,
                'created_time': stat.st_ctime,
                'modified_time': stat.st_mtime,
                'file_extension': file_path.suffix.lower()
            })
            
            # 内容统计
            metadata.update({
                'content_length': len(content),
                'word_count': len(content.split()) if content else 0,
                'line_count': content.count('\n') + 1 if content else 0,
                'char_count': len(content.replace(' ', '')) if content else 0
            })
            
            # 语言检测（简单的中英文检测）
            if content:
                chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', content))
                english_chars = len(re.findall(r'[a-zA-Z]', content))
                
                if chinese_chars > english_chars:
                    metadata['primary_language'] = 'chinese'
                elif english_chars > 0:
                    metadata['primary_language'] = 'english'
                else:
                    metadata['primary_language'] = 'unknown'
                
                metadata['chinese_char_count'] = chinese_chars
                metadata['english_char_count'] = english_chars
            
        except Exception as e:
            self.logger.warning(f"提取元数据失败: {e}")
        
        return metadata
    
    def validate_file(self, file_path: str, max_size: int = 100 * 1024 * 1024) -> bool:
        """
        验证文件是否可以处理
        
        Args:
            file_path: 文件路径
            max_size: 最大文件大小（字节）
            
        Returns:
            bool: 是否可以处理
        """
        try:
            file_path = Path(file_path)
            
            # 检查文件是否存在
            if not file_path.exists():
                return False
            
            # 检查文件大小
            if file_path.stat().st_size > max_size:
                self.logger.warning(f"文件过大: {file_path}, 大小: {file_path.stat().st_size}")
                return False
            
            # 检查文件格式
            if file_path.suffix.lower() not in self.supported_formats:
                return False
            
            return True
            
        except Exception as e:
            self.logger.error(f"文件验证失败: {e}")
            return False
    
    def get_supported_formats(self) -> list:
        """获取支持的文件格式列表"""
        return list(self.supported_formats.keys())